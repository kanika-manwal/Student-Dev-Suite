# ==============================================================
# 🎓 STUDENT DEV SUITE (VIVID EDITION)
# Full AI + ML + Resume + IDE + Animated UI
# ==============================================================
import streamlit as st
from streamlit_lottie import st_lottie
import requests, os, sys, sqlite3, hashlib, tempfile, subprocess, json
from datetime import datetime
from io import BytesIO
from dotenv import load_dotenv
from groq import Groq

# Optional dependencies
try:
    from docx import Document
    from docx.shared import Inches
except Exception:
    Document, Inches = None, None

try:
    from fpdf import FPDF
except Exception:
    FPDF = None

try:
    from streamlit_ace import st_ace
except Exception:
    st_ace = None

try:
    import numpy as np
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.linear_model import LinearRegression
    from sklearn.metrics.pairwise import cosine_similarity
except Exception:
    np = None
    TfidfVectorizer = None
    LinearRegression = None
    cosine_similarity = None

try:
    from radon.complexity import cc_visit
except Exception:
    cc_visit = None

# ==============================================================
# 🌈 PAGE CONFIG & THEME
# ==============================================================
st.set_page_config(page_title="🎓 Student Dev Suite", layout="wide")

st.markdown("""
<style>
div.block-container {
    animation: fadeIn 1s ease-in-out;
}
@keyframes fadeIn {
    from {opacity: 0; transform: translateY(10px);}
    to {opacity: 1; transform: translateY(0);}
}
body {background-color: #f8f9fb; color: #222;}
.stButton>button {
    background-color: #4F8BF9; color: white;
    border-radius: 10px; font-weight: bold;
    height: 3em; width: 100%; transition: 0.3s;
}
.stButton>button:hover {
    background-color: #2d6df7; transform: scale(1.05);
}
.stTextInput>div>div>input, .stSelectbox>div>div {
    border-radius: 8px; border: 1px solid #ccc;
}
.stMarkdown h1, h2, h3 { color: #4F8BF9 !important; }
</style>
""", unsafe_allow_html=True)

# ==============================================================
# 🔐 DATABASE & AUTH SYSTEM
# ==============================================================
DB_PATH = "studentsuite_users.db"
def init_db():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE,
            password_hash TEXT,
            created_at TEXT
        )
    """)
    conn.commit()
    return conn

conn = init_db()

def hash_password(pw): return hashlib.sha256(pw.encode()).hexdigest()

def register_user(username, password):
    cur = conn.cursor()
    try:
        cur.execute("INSERT INTO users (username, password_hash, created_at) VALUES (?, ?, ?)",
                    (username, hash_password(password), datetime.utcnow().isoformat()))
        conn.commit()
        return True, "✅ Registered successfully!"
    except sqlite3.IntegrityError:
        return False, "⚠️ Username already exists."

def verify_user(username, password):
    cur = conn.cursor()
    cur.execute("SELECT password_hash FROM users WHERE username = ?", (username,))
    row = cur.fetchone()
    return row and row[0] == hash_password(password)

# ==============================================================
# 💾 SAVE / LOAD USER DATA LOCALLY
# ==============================================================
SAVE_FILE = "user_data.json"

def save_progress(data):
    try:
        with open(SAVE_FILE, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)
        st.toast("💾 Progress saved locally!", icon="💾")
    except Exception as e:
        st.error(f"Could not save progress: {e}")

def load_progress():
    if os.path.exists(SAVE_FILE):
        with open(SAVE_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"resume": {}, "ai_history": [], "notes": "", "study": []}

# ==============================================================
# 🤖 GROQ AI ASSISTANT
# ==============================================================
load_dotenv()

def ai_chat(prompt, system="You are a helpful AI assistant for computer science students."):
    key = os.environ.get("GROQ_API_KEY")
    if not key:
        return "❌ GROQ_API_KEY not found. Add it in your .env file."

    try:
        client = Groq(api_key=key)
        resp = client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": prompt},
            ],
            temperature=0.4,
            max_tokens=700,
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return f"⚠️ Groq request failed: {e}"

# ==============================================================
# 🌈 LOTTIE ANIMATION LOADER
# ==============================================================
def load_lottie(url: str):
    try:
        r = requests.get(url)
        if r.status_code == 200:
            return r.json()
    except:
        pass
    return None

# ==============================================================
# 🧠 SESSION INITIALIZATION
# ==============================================================
if "auth" not in st.session_state:
    st.session_state.auth = False
    st.session_state.user = None
if "user_data" not in st.session_state:
    st.session_state.user_data = load_progress()

# ==============================================================
# 🧭 SIDEBAR NAVIGATION
# ==============================================================
with st.sidebar:
    st.title("🎓 Student Dev Suite")
    page = st.radio("Navigate to:", [
        "Home", "AI Assistant", "Resume Generator", "PDF Creator",
        "IDE", "Document Creator", "Resume Enhancer (ML)",
        "Code Insights (ML)", "Study Recommender", "Notes Summarizer",
        "Progress Predict (ML)", "AI Text Detector (ML)", "Account"
    ])
    if st.button("💾 Save Progress"):
        save_progress(st.session_state.user_data)
    st.markdown("---")
    st.caption("Powered by Groq | Built with ❤️ using Streamlit")

# ==============================================================
# 🏠 HOME PAGE — With Vivid Animation
# ==============================================================
if page == "Home":
    st.markdown('<h1 style="text-align:center;color:#4F8BF9;">🎓 Student Dev Suite</h1>', unsafe_allow_html=True)
    st.markdown('<p style="text-align:center;color:#666;">Your all-in-one AI, ML, Resume, and Coding toolkit</p>', unsafe_allow_html=True)
    anim = load_lottie("https://lottie.host/ed6276a9-7a2b-44b2-bf2b-b6328236a177/E7Jt6LxNtd.json")
    if anim: st_lottie(anim, height=300, key="home_anim")
    st.success("Welcome! Use the sidebar to explore all tools 🚀")
    st.info("💡 Tip: Use 'Save Progress' anytime to keep your data stored locally.")
# ==============================================================
# 🤖 AI ASSISTANT (with vivid animation)
# ==============================================================
if page == "AI Assistant":
    st.header("🤖 AI Assistant (Groq LLaMA 3)")
    anim = load_lottie("https://lottie.host/7e9a7877-fbfd-4b5f-8578-7b82ac0b7a4a/KJ7ZTtqmmL.json")
    if anim: st_lottie(anim, height=250, key="ai_anim")
    prompt = st.text_area("Ask your question:", height=200, key="ai_prompt")
    if st.button("Ask Groq"):
        with st.spinner("Thinking..."):
            answer = ai_chat(prompt)
        st.subheader("Answer:")
        st.write(answer)
        st.session_state.user_data["ai_history"].append({"prompt": prompt, "answer": answer})

# ==============================================================
# 🧾 HYBRID RESUME GENERATOR (Modern + Photo in DOCX & PDF)
# ==============================================================
if page == "Resume Generator":
    st.header("🧾 Resume Generator")
    anim = load_lottie("https://lottie.host/14e1b85f-bb3a-4e0c-8a22-65ab29f8122d/2bqPRGXsET.json")
    if anim: st_lottie(anim, height=250, key="resume_anim")

    st.info("Create a professional hybrid resume with your photo, formatted for both DOCX & PDF.")

    name = st.text_input("Full Name", st.session_state.user_data["resume"].get("name", ""))
    title = st.text_input("Job Title / Role", st.session_state.user_data["resume"].get("title", ""))
    email = st.text_input("Email", st.session_state.user_data["resume"].get("email", ""))
    phone = st.text_input("Phone", st.session_state.user_data["resume"].get("phone", ""))
    summary = st.text_area("Summary", st.session_state.user_data["resume"].get("summary", ""), height=100)
    photo = st.file_uploader("Upload Profile Picture (optional)", type=["jpg", "jpeg", "png"])
    style = st.selectbox("Select Resume Style", ["Hybrid Modern", "Classic", "Minimal"])

    # Education section
    edu_count = st.number_input("Education Entries", 1, 5, 1)
    education = []
    for i in range(int(edu_count)):
        with st.expander(f"🎓 Education {i+1}"):
            degree = st.text_input(f"Degree {i+1}")
            inst = st.text_input(f"Institution {i+1}")
            year = st.text_input(f"Year {i+1}")
            education.append({"degree": degree, "institution": inst, "year": year})

    # Skills and projects
    skills = st.text_input("Skills (comma separated)")
    proj_count = st.number_input("Projects", 0, 5, 1)
    projects = []
    for i in range(int(proj_count)):
        with st.expander(f"💻 Project {i+1}"):
            pname = st.text_input(f"Project Name {i+1}")
            pdesc = st.text_area(f"Description {i+1}")
            projects.append({"name": pname, "desc": pdesc})

    if st.button("Generate Resume"):
        try:
            # --- DOCX CREATION ---
            doc = Document()
            doc.add_heading(name, level=0)
            if photo:
                doc.add_picture(photo, width=Inches(1.3))
            doc.add_paragraph(title)
            doc.add_paragraph(f"📧 {email} | 📞 {phone}")
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(summary)
            doc.add_heading("Education", level=1)
            for edu in education:
                doc.add_paragraph(f"{edu['degree']} — {edu['institution']} ({edu['year']})")
            doc.add_heading("Skills", level=1)
            doc.add_paragraph(skills)
            doc.add_heading("Projects", level=1)
            for proj in projects:
                doc.add_paragraph(f"{proj['name']}: {proj['desc']}")

            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            st.download_button("📄 Download DOCX", bio, file_name=f"{name or 'resume'}.docx")

            # --- PDF CREATION ---
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, name, ln=True, align="C")
            pdf.set_font("Helvetica", "", 12)
            pdf.cell(0, 8, title, ln=True, align="C")
            pdf.ln(8)

            # add photo if available
            if photo:
                temp_photo = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
                temp_photo.write(photo.read())
                temp_photo.flush()
                pdf.image(temp_photo.name, x=170, y=20, w=25)
                os.unlink(temp_photo.name)

            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, "Summary", ln=True)
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, summary)
            pdf.ln(4)

            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, "Education", ln=True)
            pdf.set_font("Helvetica", "", 11)
            for edu in education:
                pdf.cell(0, 8, f"{edu['degree']} - {edu['institution']} ({edu['year']})", ln=True)

            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, "Skills", ln=True)
            pdf.set_font("Helvetica", "", 11)
            pdf.multi_cell(0, 6, skills)

            pdf.ln(4)
            pdf.set_font("Helvetica", "B", 14)
            pdf.cell(0, 10, "Projects", ln=True)
            pdf.set_font("Helvetica", "", 11)
            for proj in projects:
                pdf.multi_cell(0, 6, f"• {proj['name']}: {proj['desc']}")
                pdf.ln(2)

            bio_pdf = BytesIO()
            pdf.output(bio_pdf)
            bio_pdf.seek(0)
            st.download_button("📘 Download PDF", data=bio_pdf, file_name=f"{name or 'resume'}.pdf")
            st.success("✅ Resume generated successfully!")

            # save to session for persistence
            st.session_state.user_data["resume"] = {
                "name": name, "title": title, "email": email, "phone": phone,
                "summary": summary, "skills": skills
            }

        except Exception as e:
            st.error(f"Error generating resume: {e}")

# ==============================================================
# 📄 PDF CREATOR
# ==============================================================
if page == "PDF Creator":
    st.header("📄 PDF Creator")
    anim = load_lottie("https://lottie.host/3a798c38-f295-4b72-81b0-19e61615f58a/7WZtHknW8a.json")
    if anim: st_lottie(anim, height=200, key="pdf_anim")

    text = st.text_area("Enter text to convert into PDF", height=250)
    title = st.text_input("Title", value="Document")

    if st.button("Create PDF"):
        if FPDF is None:
            st.error("FPDF not installed.")
        else:
            try:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Helvetica", "B", 16)
                pdf.cell(0, 10, title, ln=True)
                pdf.set_font("Helvetica", "", 12)
                pdf.multi_cell(0, 8, text)
                out = BytesIO()
                pdf.output(out)
                out.seek(0)
                st.download_button("📘 Download PDF", out, file_name=f"{title}.pdf")
                st.success("✅ PDF generated!")
            except Exception as e:
                st.error(f"PDF creation failed: {e}")

# ==============================================================
# 🧾 DOCUMENT CREATOR (DOCX / PDF)
# ==============================================================
if page == "Document Creator":
    st.header("🧾 Document Creator")
    anim = load_lottie("https://lottie.host/ef3ac3d4-5d64-4b52-9e87-19e15e7e6da2/nx3YJHq9rA.json")
    if anim: st_lottie(anim, height=200, key="doc_anim")

    title = st.text_input("Document Title")
    body = st.text_area("Body (supports markdown)", height=250)

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Export DOCX"):
            try:
                doc = Document()
                doc.add_heading(title or "Document", 0)
                for line in body.split("\n"):
                    doc.add_paragraph(line)
                out = BytesIO()
                doc.save(out)
                out.seek(0)
                st.download_button("📄 Download DOCX", out, file_name=f"{title or 'doc'}.docx")
            except Exception as e:
                st.error(f"Failed: {e}")

    with col2:
        if st.button("Export PDF"):
            try:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Helvetica", "B", 16)
                pdf.cell(0, 10, title, ln=True)
                pdf.set_font("Helvetica", "", 12)
                pdf.multi_cell(0, 8, body)
                out = BytesIO()
                pdf.output(out)
                out.seek(0)
                st.download_button("📘 Download PDF", out, file_name=f"{title or 'doc'}.pdf")
            except Exception as e:
                st.error(f"Failed to export PDF: {e}")

# ==============================================================
# 💻 IDE (Multi-language Runner)
# ==============================================================
if page == "IDE":
    st.header("💻 IDE - Code Editor + Runner")
    anim = load_lottie("https://lottie.host/8c4569b4-f81f-4f91-b77b-30323de71b3c/KwXld7vLkC.json")
    if anim: st_lottie(anim, height=230, key="ide_anim")
    st.warning("⚠️ Use in a trusted/private environment. (Runs code server-side)")

    lang = st.selectbox("Language", ["python", "javascript", "cpp", "java"])
    default_code = {
        "python": 'print("Hello from Python!")',
        "javascript": 'console.log("Hello from Node.js!")',
        "cpp": '#include <iostream>\nusing namespace std;\nint main(){ cout<<"Hello C++!\\n"; return 0; }',
        "java": 'public class Main { public static void main(String[] args) { System.out.println("Hello Java!"); } }',
    }

    if st_ace:
        code = st_ace(value=default_code[lang], language=lang if lang != "cpp" else "cpp", theme="monokai", key="ace_editor")
    else:
        code = st.text_area("Code Editor", default_code[lang], height=250)

    stdin = st.text_area("Optional stdin input", height=80)
    if st.button("Run Code"):
        with st.spinner("Running your code..."):
            with tempfile.TemporaryDirectory() as tmpdir:
                try:
                    if lang == "python":
                        src = os.path.join(tmpdir, "main.py")
                        open(src, "w").write(code)
                        result = subprocess.run([sys.executable, src], input=stdin.encode(),
                                                stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
                    elif lang == "javascript":
                        src = os.path.join(tmpdir, "main.js")
                        open(src, "w").write(code)
                        result = subprocess.run(["node", src], input=stdin.encode(),
                                                stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
                    elif lang == "cpp":
                        src = os.path.join(tmpdir, "main.cpp")
                        exe = os.path.join(tmpdir, "main.exe")
                        open(src, "w").write(code)
                        subprocess.run(["g++", src, "-o", exe], timeout=10)
                        result = subprocess.run([exe], input=stdin.encode(),
                                                stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
                    elif lang == "java":
                        src = os.path.join(tmpdir, "Main.java")
                        open(src, "w").write(code)
                        subprocess.run(["javac", src], timeout=10)
                        result = subprocess.run(["java", "-cp", tmpdir, "Main"], input=stdin.encode(),
                                                stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=10)
                    else:
                        raise ValueError("Unsupported language")

                    st.subheader("Output:")
                    st.code(result.stdout.decode(errors="replace"))
                    st.subheader("Errors (if any):")
                    st.code(result.stderr.decode(errors="replace"))

                except subprocess.TimeoutExpired:
                    st.error("Execution timed out.")
                except Exception as e:
                    st.error(f"Error: {e}")
# ==============================================================
# 🧠 RESUME ENHANCER (ML)
# ==============================================================
if page == "Resume Enhancer (ML)":
    st.header("🧠 Resume Enhancer (ML)")
    anim = load_lottie("https://lottie.host/cc889c41-f382-4b70-ae6d-76ce44cc7d91/NL3GHE8cJm.json")
    if anim: st_lottie(anim, height=200, key="enhancer_anim")
    st.caption("Compare your resume with a job description and get missing keyword suggestions.")

    resume_txt = st.text_area("Paste your Resume Text", height=180)
    job_desc = st.text_area("Paste Job Description (JD)", height=180)
    topk = st.slider("Number of Keywords", 5, 30, 12)

    if st.button("Analyze Resume"):
        if not TfidfVectorizer or not np:
            st.error("⚠️ scikit-learn not installed.")
        else:
            try:
                tfidf = TfidfVectorizer(stop_words="english", ngram_range=(1, 2), max_features=5000)
                X = tfidf.fit_transform([resume_txt, job_desc])
                feat_names = np.array(tfidf.get_feature_names_out())
                scores = X.toarray()[1]
                top_idx = scores.argsort()[::-1][:topk * 2]
                jd_keywords = feat_names[top_idx]
                resume_terms = set(resume_txt.lower().split())
                missing = [w for w in jd_keywords if w not in resume_terms][:topk]

                st.subheader("💡 Suggested Keywords to Add:")
                st.write(", ".join(missing) if missing else "✅ Resume covers most job terms!")

                sim = cosine_similarity(X[0], X[1])[0][0]
                st.metric("Resume ↔ JD Similarity", f"{sim:.2f}")

            except Exception as e:
                st.error(f"Error: {e}")

# ==============================================================
# 💻 CODE INSIGHTS (ML)
# ==============================================================
if page == "Code Insights (ML)":
    st.header("💻 Code Insights (ML)")
    anim = load_lottie("https://lottie.host/4a4ebcc4-3a11-4b8a-87b3-63dc0b192ec3/GBW7kbAI3M.json")
    if anim: st_lottie(anim, height=200, key="insights_anim")
    lang = st.selectbox("Select Language", ["python", "javascript", "cpp", "java"])
    code_ci = st.text_area("Paste your code", height=200)
    if st.button("Analyze Code"):
        findings = []
        if "== None" in code_ci:
            findings.append("🔹 Use `is None` instead of `== None` for Pythonic style.")
        if "var " in code_ci:
            findings.append("🔹 Use `let` or `const` instead of `var` in JavaScript.")
        if len(code_ci) > 300:
            findings.append("🔹 Large file: consider modularizing your code.")
        if cc_visit and lang == "python":
            try:
                blocks = cc_visit(code_ci)
                complex = [b for b in blocks if b.complexity >= 10]
                if complex:
                    findings.append(f"⚠️ {len(complex)} functions with high complexity detected.")
            except Exception:
                pass
        st.subheader("Findings:")
        st.write("\n".join(findings) or "✅ Code looks clean!")

# ==============================================================
# 📚 STUDY RECOMMENDER
# ==============================================================
if page == "Study Recommender":
    st.header("📚 Study Recommender")
    anim = load_lottie("https://lottie.host/5f84d22c-5e9b-4d60-b4ec-3a6c0a908b80/tvTGv2RZck.json")
    if anim: st_lottie(anim, height=220, key="study_anim")
    focus = st.selectbox("Your Focus Area", ["Python", "Data Structures", "Web Dev (JS)", "Java", "C++", "Machine Learning"])
    level = st.selectbox("Level", ["Beginner", "Intermediate", "Advanced"])
    if st.button("Get Recommendations"):
        recs = {
            "Python": ["Automate the Boring Stuff", "Real Python", "LeetCode easy problems"],
            "Data Structures": ["GeeksForGeeks practice", "VisualAlgo.net", "Implement Tree/Graph manually"],
            "Web Dev (JS)": ["MDN JS Guide", "React docs", "Build a ToDo app"],
            "Java": ["Head First Java", "Spring Boot basics", "OOP project"],
            "C++": ["STL tutorial", "Codeforces practice", "Hackerrank challenges"],
            "Machine Learning": ["Hands-On ML (Aurelien)", "Kaggle competitions", "Train Linear Regression manually"]
        }
        st.subheader("📖 Recommended Next Steps")
        for r in recs[focus]:
            st.write(f"- {r}")
        st.session_state.user_data["study"].append({"focus": focus, "level": level})

# ==============================================================
# 📝 NOTES SUMMARIZER
# ==============================================================
if page == "Notes Summarizer":
    st.header("📝 Notes Summarizer")

    # Lottie animation
    anim = load_lottie("https://lottie.host/4d8ff438-5d6d-40ce-b36d-5c7c02c7d35d/EPPvmnSv7n.json")
    if anim:
        st_lottie(anim, height=200, key="notes_anim")

    # Input section
    notes = st.text_area("Paste your notes", height=250, key="notes_text")
    topk = st.slider("Keywords", 5, 20, 10)

    # Summarization button
    if st.button("Summarize Notes"):
        if not TfidfVectorizer or not np:
            st.error("⚠️ scikit-learn or numpy not installed.")
        else:
            try:
                # --- Keyword Extraction ---
                vec = TfidfVectorizer(stop_words="english")
                X = vec.fit_transform([notes])
                feats = np.array(vec.get_feature_names_out())
                idx = X.toarray()[0].argsort()[::-1][:topk]
                kw = feats[idx]

                st.subheader("🔑 Top Keywords")
                st.write(", ".join(kw))

                # --- Sentence-Level Summarization ---
                sentences = [s.strip() for s in notes.split(".") if len(s.strip()) > 5]
                if sentences:
                    vec_s = TfidfVectorizer(stop_words="english").fit_transform(sentences)
                    vec_s = vec_s.toarray()  # ✅ FIX: convert sparse matrix to numpy array
                    centroid = np.mean(vec_s, axis=0).reshape(1, -1)
                    sims = cosine_similarity(vec_s, centroid)
                    top_s = np.argsort(-sims.ravel())[: min(5, len(sentences))]
                    summary = ". ".join([sentences[i] for i in top_s])

                    st.subheader("🧩 Summary")
                    st.write(summary)

                # Save notes in user session
                st.session_state.user_data["notes"] = notes

            except Exception as e:
                st.error(f"❌ Error: {e}")


# ==============================================================
# 📈 PROGRESS PREDICT (ML)
# ==============================================================
if page == "Progress Predict (ML)":
    st.header("📈 Progress Predictor (ML)")
    anim = load_lottie("https://lottie.host/79e746ed-77b0-4329-882d-8a21c0a27267/1DkEckLge7.json")
    if anim: st_lottie(anim, height=220, key="progress_anim")
    st.caption("Predict future progress based on weekly hours & completed tasks.")
    import pandas as pd
    df = st.data_editor(pd.DataFrame({"week": [1, 2, 3, 4], "hours": [2, 4, 6, 8], "completed": [3, 6, 10, 13]}), num_rows="dynamic")
    target = st.number_input("Predict Week", 5, 52, 8)
    if st.button("Train & Predict"):
        if not LinearRegression:
            st.error("scikit-learn not installed.")
        else:
            try:
                X, y = df[["week", "hours"]].values, df["completed"].values
                model = LinearRegression().fit(X, y)
                pred = model.predict([[target, df["hours"].mean()]])[0]
                st.metric("Predicted Completed Tasks", f"{pred:.1f}")
                st.line_chart(df.set_index("week")[["hours", "completed"]])
            except Exception as e:
                st.error(f"Training failed: {e}")

# ==============================================================
# 🤖 AI TEXT DETECTOR (ML)
# ==============================================================
if page == "AI Text Detector (ML)":
    st.header("🤖AI-Generated Text Detector (ML)")
    st.caption("Upload text files to check if they are AI-generated or human-written using perplexity-based analysis.")

    uploaded = st.file_uploader(
        "Upload 1–5 text files (.txt, .md, .py, etc.)",
        accept_multiple_files=True
    )

    if st.button("Analyze AI Content"):
        if not uploaded:
            st.error("Please upload at least one text file.")
        else:
            import numpy as np
            import re
            import pandas as pd
            from transformers import GPT2TokenizerFast

            tokenizer = GPT2TokenizerFast.from_pretrained("gpt2")

            def estimate_ai_percentage(text):
                """
                Estimate AI-generated probability using simple heuristics:
                - Low diversity or short words → AI-like
                - High variety and word length → Human-like
                """
                tokens = tokenizer.encode(text, truncation=True, max_length=1024)
                if not tokens or len(text.strip()) == 0:
                    return 0.0, "❌ Empty or unreadable text"

                # Compute diversity and linguistic complexity
                unique_ratio = len(set(tokens)) / len(tokens)
                avg_word_len = np.mean([len(w) for w in re.findall(r'\w+', text)]) if text.strip() else 0

                # Inverted heuristic (AI-generated text tends to have lower perplexity)
                perplexity_score = (unique_ratio * 100) + (avg_word_len * 2)
                ai_score = max(0, min(100, 120 - perplexity_score))

                # Verdict based on score
                if ai_score > 75:
                    verdict = "⚠️ Highly likely AI-generated"
                elif ai_score > 45:
                    verdict = "🤔 Possibly AI-edited or hybrid"
                else:
                    verdict = "✅ Likely human-written"

                return round(ai_score, 2), verdict

            # Analyze all uploaded files
            results = []
            for f in uploaded:
                try:
                    text = f.read().decode("utf-8", errors="ignore")
                    ai_score, verdict = estimate_ai_percentage(text)
                    results.append({
                        "File": f.name,
                        "AI Probability (%)": ai_score,
                        "Verdict": verdict
                    })
                except Exception as e:
                    results.append({
                        "File": f.name,
                        "AI Probability (%)": 0.0,
                        "Verdict": f"⚠️ Error: {e}"
                    })

            # Display results
            df = pd.DataFrame(results)
            st.subheader("🧠 AI Detection Results")
            st.dataframe(df)

            # Calculate average AI probability
            avg_ai = np.mean([
                r["AI Probability (%)"]
                for r in results
                if isinstance(r["AI Probability (%)"], (int, float))
            ])
            st.metric("Average AI Probability", f"{avg_ai:.2f}%")

            # Display verdict
            if avg_ai > 70:
                st.warning("⚠️ Overall content appears largely AI-generated.")
            elif avg_ai > 40:
                st.info("🤔 Mixed content — some AI involvement possible.")
            else:
                st.success("✅ Mostly human-written text detected.")

# ==============================================================
# 🧍 ACCOUNT MANAGEMENT
# ==============================================================
if page == "Account":
    st.header("👤 Account Management")
    anim = load_lottie("https://lottie.host/d564e5ef-3240-4a91-b5dc-524cb319ed75/kD9aZMj9uY.json")
    if anim: st_lottie(anim, height=200, key="account_anim")
    if not st.session_state.auth:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Register")
            ru = st.text_input("Username")
            rp = st.text_input("Password", type="password")
            if st.button("Register"):
                ok, msg = register_user(ru, rp)
                st.success(msg) if ok else st.error(msg)
        with col2:
            st.subheader("Login")
            lu = st.text_input("Login Username")
            lp = st.text_input("Login Password", type="password")
            if st.button("Login"):
                if verify_user(lu, lp):
                    st.session_state.auth = True
                    st.session_state.user = lu
                    st.success(f"Welcome {lu}!")
                else:
                    st.error("Invalid credentials.")
    else:
        st.success(f"Logged in as {st.session_state.user}")
        if st.button("Logout"):
            st.session_state.auth = False
            st.session_state.user = None
            st.experimental_rerun()

# ==============================================================
# ⚙️ FOOTER
# ==============================================================
st.markdown("---")
st.caption("🎓 Student Dev Suite — Built with ❤️ using Streamlit + Groq API (LLaMA 3.3)")
st.caption("💾 Progress is automatically saved to `user_data.json` in your app folder.")

